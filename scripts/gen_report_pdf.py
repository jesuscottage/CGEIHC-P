"""
Genera los reportes en DOCX profesional con python-docx.
Uso: python gen_report_pdf.py [es|en|both]
"""
import sys, os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.join(SCRIPT_DIR, '..')
IMG_DIR = os.path.join(PROJECT_DIR, 'active')
ENTREGABLES_DIR = os.path.join(PROJECT_DIR, 'entregables')

BLACK = RGBColor(0, 0, 0)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
CODE_BG = RGBColor(0xF5, 0xF5, 0xF5)

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading.append(shd)

def style_doc(doc):
    """Configurar estilos base del documento."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = BLACK
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for i in range(1, 5):
        h = doc.styles[f'Heading {i}']
        h.font.color.rgb = BLACK
        h.font.name = 'Arial'
        h.font.bold = True
        if i == 1:
            h.font.size = Pt(18)
        elif i == 2:
            h.font.size = Pt(15)
        elif i == 3:
            h.font.size = Pt(13)
        else:
            h.font.size = Pt(12)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

def add_para(doc, text, bold=False, size=None, align=None, italic=False, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = BLACK
    run.font.name = 'Arial'
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_code_block(doc, code, lang='glsl'):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_GRAY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.color.rgb = BLACK
    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            run.font.color.rgb = BLACK
    doc.add_paragraph()  # spacer

def add_figure(doc, img_name, caption, width=Inches(5.5)):
    img_path = os.path.join(IMG_DIR, img_name)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=width)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = BLACK
        run.font.name = 'Arial'
        cap.paragraph_format.space_after = Pt(12)

def build_cover(doc, lang='es'):
    """Portada profesional sin decir 'PORTADA'."""
    # Espaciado superior
    for _ in range(3):
        add_para(doc, '', space_after=0)

    # Logos centrados
    logo_fi = os.path.join(IMG_DIR, 'portada_img1.png')
    logo_unam = os.path.join(IMG_DIR, 'portada_img2.png')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(logo_fi):
        run = p.add_run()
        run.add_picture(logo_fi, width=Inches(1.0))
    run = p.add_run('     ')
    if os.path.exists(logo_unam):
        run = p.add_run()
        run.add_picture(logo_unam, width=Inches(1.0))

    add_para(doc, '', space_after=6)
    add_para(doc, 'Universidad Nacional Autónoma de México',
             bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, 'Facultad de Ingeniería' if lang == 'es' else 'Faculty of Engineering',
             bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, 'Computación Gráfica e Interacción Humano-Computadora' if lang == 'es'
             else 'Computer Graphics and Human-Computer Interaction',
             size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    # Línea separadora
    add_para(doc, '─' * 60, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Título del proyecto
    title = ('Nuestro Mundo — Un museo interactivo sobre el cambio climático '
             'en nuestro planeta y su afectación en el polo norte') if lang == 'es' else \
            ('Our World — An interactive museum about climate change '
             'on our planet and its impact on the North Pole')
    add_para(doc, title, bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    add_para(doc, '─' * 60, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    # Info
    gt = 'Grupo Teoría: 2  |  Equipo: 9' if lang == 'es' else 'Theory Group: 2  |  Team: 9'
    add_para(doc, gt, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    doc_label = 'Docente:' if lang == 'es' else 'Professor:'
    add_para(doc, f'{doc_label} Dr. Sergio Teodoro Vite',
             bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    al = 'Alumnos:' if lang == 'es' else 'Students:'
    add_para(doc, al, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    for name in ['Navarro Nuño Juan Pablo', 'Martínez Jiménez Israel', 'Hernández Cabañas Jesús']:
        add_para(doc, name, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    add_para(doc, '', space_after=20)
    add_para(doc, 'Semestre 2026-2', bold=True, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    doc.add_page_break()


def build_bitacora(doc, lang='es'):
    title = 'Bitácora de colaboración' if lang == 'es' else 'Collaboration Log'
    doc.add_heading(title, level=1)

    headers_es = ['Fecha', 'Acción', 'Colaborador', 'Notas']
    headers_en = ['Date', 'Action', 'Collaborator', 'Notes']
    headers = headers_es if lang == 'es' else headers_en

    rows_es = [
        ['2 abr 2026', 'Definir la idea del proyecto y elaborar bocetos del museo', 'Juan Pablo', 'Estuvimos decidiendo entre un acuario y un museo; al final elegimos el museo porque daba más juego narrativo'],
        ['5 abr 2026', 'Configurar el entorno de desarrollo (CMake, GLFW, GLAD, GLM)', 'Jesús', 'Tuve que investigar cómo usar FetchContent en CMake porque no queríamos instalar dependencias manualmente'],
        ['10 abr 2026', 'Crear la estructura del museo: suelo, paredes, techo y cámara FPS', 'Juan Pablo', 'La primera versión era un solo pasillo; después lo cambiamos a forma de T con dos corredores'],
        ['14 abr 2026', 'Implementar el shader de iluminación Blinn-Phong y Fresnel', 'Jesús', 'Nos basamos en los shaders del profesor como referencia inicial y los fuimos adaptando'],
        ['17 abr 2026', 'Descargar y preparar modelos 3D (turbina, oso polar, edificios)', 'Israel', 'Busqué modelos CC0 en Quaternius y Kenney; algunos venían sin textura y los ajustamos con color override'],
        ['20 abr 2026', 'Agregar el skybox cubemap y el plano de agua con olas procedurales', 'Juan Pablo', 'El shader de agua fue lo más difícil; probé con una sola onda y se veía artificial, al final usé tres ondas cruzadas'],
        ['23 abr 2026', 'Implementar las animaciones de los 7 módulos interactivos', 'Jesús', 'Cada módulo tiene su animación de 8 segundos; usé glm::mix para interpolar las propiedades'],
        ['26 abr 2026', 'Integrar audio con miniaudio (ambiente, narraciones, efectos)', 'Israel', 'Primero intenté con OpenAL pero daba muchos problemas con las DLLs; miniaudio fue mucho más fácil porque es un solo header'],
        ['28 abr 2026', 'Crear la galería de pinturas y el sistema de popups narrativos', 'Juan Pablo', 'Generé las 12 pinturas con un script de Python y las distribuí en los dos corredores'],
        ['30 abr 2026', 'Agregar colisiones, sistema de progreso y pantalla de cierre', 'Jesús', 'Las colisiones bidireccionales con las paredes interiores me tomaron varias iteraciones'],
        ['2 may 2026', 'Pruebas generales, corrección de escalas y preparación para la primera exposición', 'Israel', 'Revisamos que todos los módulos se activaran correctamente y que el audio no se cortara'],
        ['26 may 2026', 'Agregar point lights en las exhibiciones y reducir luz ambiental', 'Jesús', 'El profesor nos pidió que las exhibiciones se distinguieran más del pasillo; agregamos 8 luces puntuales cálidas'],
        ['27 may 2026', 'Corregir orientación de modelos, arreglar animación del auto y reset de animaciones', 'Juan Pablo', 'El oso estaba con la cabeza hacia abajo y el auto se movía de lado; corregimos las rotaciones'],
        ['28 may 2026', 'Crear el globo terráqueo procedural, arreglar el iceberg y generar el instalador', 'Jesús', 'El modelo GLB del iceberg no renderizaba, así que lo reemplacé con geometría procedural'],
    ]
    rows_en = [
        ['Apr 2, 2026', 'Define the project idea and create museum sketches', 'Juan Pablo', 'We were deciding between an aquarium and a museum; we chose the museum because it offered more narrative possibilities'],
        ['Apr 5, 2026', 'Set up the development environment (CMake, GLFW, GLAD, GLM)', 'Jesús', 'I had to research how to use FetchContent in CMake to avoid installing dependencies manually'],
        ['Apr 10, 2026', 'Build the museum structure: floor, walls, ceiling and first-person camera', 'Juan Pablo', 'The first version was a single hallway; we later changed it to a T-shape with two corridors'],
        ['Apr 14, 2026', 'Implement the Blinn-Phong and Fresnel lighting shader', 'Jesús', 'We used the professor\'s shaders as a starting reference and adapted them to our needs'],
        ['Apr 17, 2026', 'Download and prepare 3D models (turbine, polar bear, buildings)', 'Israel', 'I searched for CC0 models on Quaternius and Kenney; some came without textures so we used color overrides'],
        ['Apr 20, 2026', 'Add the skybox cubemap and the procedural wave water plane', 'Juan Pablo', 'The water shader was the hardest part; a single wave looked artificial, so I ended up using three crossed waves'],
        ['Apr 23, 2026', 'Implement animations for the 7 interactive modules', 'Jesús', 'Each module has an 8-second animation; I used glm::mix to interpolate properties'],
        ['Apr 26, 2026', 'Integrate audio with miniaudio (ambient, narrations, effects)', 'Israel', 'I first tried OpenAL but it caused too many DLL issues; miniaudio was much easier since it is a single header'],
        ['Apr 28, 2026', 'Create the painting gallery and the narrative popup system', 'Juan Pablo', 'I generated the 12 paintings with a Python script and distributed them across both corridors'],
        ['Apr 30, 2026', 'Add collisions, progress tracking and closing screen', 'Jesús', 'The bidirectional collisions with interior walls took several iterations to get right'],
        ['May 2, 2026', 'General testing, scale corrections and preparation for the first presentation', 'Israel', 'We verified that all modules activated correctly and that audio did not cut off'],
        ['May 26, 2026', 'Add point lights on exhibitions and reduce ambient lighting', 'Jesús', 'The professor asked us to make the exhibitions stand out more; we added 8 warm point lights'],
        ['May 27, 2026', 'Fix model orientation, repair electric car animation and add animation reset', 'Juan Pablo', 'The bear was upside down and the car moved sideways; we fixed the rotations'],
        ['May 28, 2026', 'Create the procedural globe, fix the iceberg and generate the installer', 'Jesús', 'The iceberg GLB model did not render, so I replaced it with procedural geometry'],
    ]
    rows = rows_es if lang == 'es' else rows_en
    add_table(doc, headers, rows)
    doc.add_page_break()


def build_body(doc, lang='es'):
    """Cuerpo del reporte (resumen → referencias)."""
    # ── RESUMEN ──
    doc.add_heading('Resumen' if lang == 'es' else 'Abstract', level=1)
    if lang == 'es':
        add_para(doc, 'Este trabajo presenta el desarrollo de un sistema gráfico interactivo 3D en tiempo real que simula un museo virtual educativo ubicado en el Polo Norte, con el propósito de comunicar las causas, consecuencias y posibles soluciones del calentamiento global. El sistema fue desarrollado en C++ con OpenGL 3.3 Core Profile e implementa técnicas de iluminación Blinn-Phong con reflexión Fresnel, environment mapping mediante cubemap, y tres tipos de animación: básica, basada en keyframes y procedural. El museo consta de siete módulos interactivos organizados en dos rutas narrativas — consecuencias y soluciones — con narración de voz contextual, interfaz gráfica informativa y un sistema de progreso que guía al usuario desde una zona de bienvenida hasta una reflexión final.')
        p = add_para(doc, '')
        run = p.add_run('Palabras clave: ')
        run.bold = True
        run.font.color.rgb = BLACK
        run = p.add_run('computación gráfica, OpenGL, museo virtual, calentamiento global, iluminación Blinn-Phong, animación procedural')
        run.font.color.rgb = BLACK
    else:
        add_para(doc, 'This work presents the development of a real-time 3D interactive graphics system that simulates an educational virtual museum located at the North Pole, designed to communicate the causes, consequences and potential solutions of global warming. The system was developed in C++ with OpenGL 3.3 Core Profile and implements Blinn-Phong lighting with Fresnel reflection, environment mapping through cubemaps, and three types of animation: basic, keyframe-based and procedural. The museum consists of seven interactive modules organized in two narrative routes — consequences and solutions — with contextual voice narration, an informative graphical interface, and a progress system that guides the user from a welcome zone to a final reflection.')
        p = add_para(doc, '')
        run = p.add_run('Keywords: ')
        run.bold = True
        run.font.color.rgb = BLACK
        run = p.add_run('computer graphics, OpenGL, virtual museum, global warming, Blinn-Phong lighting, procedural animation')
        run.font.color.rgb = BLACK

    doc.add_page_break()

    # ── 1. INTRODUCCIÓN ──
    doc.add_heading('1. ' + ('Introducción' if lang == 'es' else 'Introduction'), level=1)
    if lang == 'es':
        add_para(doc, 'El calentamiento global representa una de las amenazas más documentadas para los ecosistemas polares. Según el National Snow and Ice Data Center (NSIDC, 2024), el hielo ártico de más de cuatro años ha perdido el 95% de su extensión desde 1979. La Unión Internacional para la Conservación de la Naturaleza proyecta una reducción poblacional de osos polares superior al 30% para 2050 (IUCN, 2023). Estos datos evidencian la necesidad de herramientas educativas que comuniquen la magnitud del problema de forma accesible.')
        add_para(doc, 'La computación gráfica ofrece un medio para representar fenómenos complejos de manera visual e interactiva. Según Angel y Shreiner (2011), los sistemas gráficos interactivos permiten a los usuarios explorar escenarios que serían inaccesibles en la realidad, facilitando la comprensión de procesos abstractos. En el ámbito educativo, los museos virtuales han demostrado ser efectivos para la divulgación científica al combinar narrativa, visualización y participación activa del usuario (Sylaiou et al., 2010).')
        add_para(doc, 'El presente proyecto tiene como objetivo diseñar e implementar un sistema gráfico 3D interactivo que simule un museo en el Polo Norte, aplicando las técnicas fundamentales de computación gráfica estudiadas durante el curso: pipeline gráfico, modelado geométrico, iluminación, texturizado, animación e interacción humano-computadora.')
    else:
        add_para(doc, 'Global warming represents one of the most documented threats to polar ecosystems. According to the National Snow and Ice Data Center (NSIDC, 2024), Arctic ice older than four years has lost 95% of its extent since 1979. The International Union for Conservation of Nature projects a polar bear population reduction exceeding 30% by 2050 (IUCN, 2023). These data highlight the need for educational tools that communicate the magnitude of the problem in an accessible manner.')
        add_para(doc, 'Computer graphics provides a medium for representing complex phenomena visually and interactively. According to Angel and Shreiner (2011), interactive graphics systems allow users to explore scenarios that would be inaccessible in reality, facilitating comprehension of abstract processes. In the educational domain, virtual museums have proven effective for scientific communication by combining narrative, visualization and active user participation (Sylaiou et al., 2010).')
        add_para(doc, 'The objective of this project is to design and implement a 3D interactive graphics system that simulates a museum at the North Pole, applying the fundamental computer graphics techniques studied during the course: graphics pipeline, geometric modeling, lighting, texturing, animation and human-computer interaction.')

    # ── 2. ESTADO DEL ARTE ──
    doc.add_heading('2. ' + ('Estado del arte' if lang == 'es' else 'State of the Art'), level=1)
    if lang == 'es':
        add_para(doc, 'Los museos virtuales 3D han sido objeto de desarrollo desde las primeras aplicaciones de realidad virtual en la década de 1990. Proyectos como el Virtual Museum of Iraq (CNR-ISTI, 2009) demostraron que la navegación en primera persona a través de espacios museísticos reconstruidos digitalmente permite una experiencia educativa comparable a la presencial.')
        add_para(doc, 'En el contexto de OpenGL educativo, De Vries (2020) sistematizó las técnicas fundamentales de renderizado moderno en su recurso LearnOpenGL, que abarca desde la configuración del pipeline gráfico hasta la implementación de iluminación PBR y environment mapping. Este recurso constituye la referencia técnica principal del presente proyecto.')
        add_para(doc, 'Respecto a la temática ambiental, la Agencia Internacional de Energía (IEA, 2023) reporta que las energías renovables alcanzaron casi el 30% de la generación eléctrica mundial, mientras que el Programa de las Naciones Unidas para el Medio Ambiente (UNEP, 2023) advierte que la trayectoria actual de emisiones conduce a un calentamiento de 2.5 a 2.9 °C para 2100, por encima del objetivo de 1.5 °C del Acuerdo de París.')
    else:
        add_para(doc, 'Virtual 3D museums have been the subject of development since the first virtual reality applications in the 1990s. Projects such as the Virtual Museum of Iraq (CNR-ISTI, 2009) demonstrated that first-person navigation through digitally reconstructed museum spaces enables an educational experience comparable to in-person visits.')
        add_para(doc, 'In the context of educational OpenGL, De Vries (2020) systematized the fundamental techniques of modern rendering in his resource LearnOpenGL, covering topics from graphics pipeline configuration to PBR lighting and environment mapping implementation. This resource constitutes the primary technical reference of the present project.')
        add_para(doc, 'Regarding the environmental theme, the International Energy Agency (IEA, 2023) reports that renewable energies reached nearly 30% of global electricity generation, while the United Nations Environment Programme (UNEP, 2023) warns that the current emissions trajectory leads to a warming of 2.5 to 2.9 °C by 2100, above the Paris Agreement target of 1.5 °C.')

    # ── 3. METODOLOGÍA ──
    doc.add_heading('3. ' + ('Metodología' if lang == 'es' else 'Methodology'), level=1)

    # 3.1
    doc.add_heading('3.1 ' + ('Arquitectura del sistema' if lang == 'es' else 'System Architecture'), level=2)
    if lang == 'es':
        add_para(doc, 'El sistema se organiza en módulos independientes que se comunican a través de un bucle principal (game loop) con las fases de entrada, actualización y renderizado. La estructura del código se divide en tres capas: Core (Window, Input, Time, CameraFPS), Graphics (Shader, Mesh, Texture, Model, Skybox) y Scene (Museum, ModuleScene, DecoScene, GalleryScene, SnowSystem).')
    else:
        add_para(doc, 'The system is organized in independent modules communicating through a main loop (game loop) with input, update and render phases. The code structure is divided into three layers: Core (Window, Input, Time, CameraFPS), Graphics (Shader, Mesh, Texture, Model, Skybox) and Scene (Museum, ModuleScene, DecoScene, GalleryScene, SnowSystem).')

    # 3.2
    doc.add_heading('3.2 ' + ('Pipeline gráfico' if lang == 'es' else 'Graphics Pipeline'), level=2)
    if lang == 'es':
        add_para(doc, 'El pipeline sigue la secuencia estándar de OpenGL 3.3 Core Profile: modelado geométrico (procedural e importado vía Assimp), transformaciones MVP con GLM, texturizado 2D y cubemap, iluminación Blinn-Phong con Fresnel Schlick, y renderizado con depth test, face culling y niebla exponencial cuadrática.')
    else:
        add_para(doc, 'The pipeline follows the standard OpenGL 3.3 Core Profile sequence: geometric modeling (procedural and imported via Assimp), MVP transformations with GLM, 2D and cubemap texturing, Blinn-Phong lighting with Fresnel Schlick, and rendering with depth test, face culling and quadratic exponential fog.')

    # 3.3
    doc.add_heading('3.3 ' + ('Bibliotecas y herramientas' if lang == 'es' else 'Libraries and Tools'), level=2)
    lib_headers = ['Biblioteca' if lang == 'es' else 'Library', 'Versión' if lang == 'es' else 'Version',
                   'Propósito' if lang == 'es' else 'Purpose', 'URL']
    lib_rows = [
        ['OpenGL', '3.3 Core', 'API gráfica principal' if lang == 'es' else 'Main graphics API', 'khronos.org/opengl'],
        ['GLSL', '3.30', 'Shaders para GPU', '—'],
        ['GLFW', '3.4', ('Ventana, contexto GL, entrada' if lang == 'es' else 'Window, GL context, input'), 'glfw.org'],
        ['GLAD', '2.x', ('Cargador de extensiones OpenGL' if lang == 'es' else 'OpenGL extension loader'), 'github.com/Dav1dde/glad'],
        ['GLM', '1.0.1', ('Matemáticas vectoriales/matriciales' if lang == 'es' else 'Vector/matrix math'), 'github.com/g-truc/glm'],
        ['Assimp', '5.4.3', ('Importación de modelos 3D' if lang == 'es' else '3D model import'), 'assimp.org'],
        ['Dear ImGui', '1.91.6', ('Interfaz gráfica inmediata' if lang == 'es' else 'Immediate mode GUI'), 'github.com/ocornut/imgui'],
        ['stb_image', '2.x', ('Carga de texturas PNG/JPG' if lang == 'es' else 'PNG/JPG texture loading'), 'github.com/nothings/stb'],
        ['stb_image_write', '1.x', ('Escritura de PNG' if lang == 'es' else 'PNG writing'), 'github.com/nothings/stb'],
        ['miniaudio', '0.11+', ('Motor de audio (MP3, WAV, OGG)' if lang == 'es' else 'Audio engine (MP3, WAV, OGG)'), 'miniaud.io'],
        ['CMake', '3.20+', ('Sistema de compilación' if lang == 'es' else 'Build system'), 'cmake.org'],
    ]
    add_table(doc, lib_headers, lib_rows)

    # 3.4 Iluminación
    doc.add_heading('3.4 ' + ('Técnicas de iluminación' if lang == 'es' else 'Lighting Techniques'), level=2)
    if lang == 'es':
        add_para(doc, 'Se implementaron cinco técnicas de iluminación en el fragment shader:')
        add_para(doc, 'Blinn-Phong: contribución difusa (NdotL) y especular con half-vector:', bold=True)
    else:
        add_para(doc, 'Five lighting techniques were implemented in the fragment shader:')
        add_para(doc, 'Blinn-Phong: diffuse (NdotL) and specular with half-vector:', bold=True)
    add_code_block(doc, 'float NdotL = max(dot(norm, ldir), 0.0);\nfloat NdotH = max(dot(norm, halfDir), 0.0);\nfloat spec = pow(NdotH, shininess);')

    add_para(doc, 'Fresnel Schlick:', bold=True)
    add_code_block(doc, 'vec3 fresnelSchlick(float cosTheta, vec3 F0) {\n    return F0 + (1.0 - F0) * pow(clamp(1.0 - cosTheta, 0.0, 1.0), 5.0);\n}')

    if lang == 'es':
        add_para(doc, 'Luz hemisférica: interpola entre color de suelo y cielo según la orientación de la normal. Point lights: hasta 8 luces puntuales con atenuación cuadrática sobre las exhibiciones. Niebla exponencial cuadrática: atenúa objetos distantes.')
    else:
        add_para(doc, 'Hemispheric light: interpolates between ground and sky color based on normal orientation. Point lights: up to 8 point lights with quadratic attenuation over exhibitions. Quadratic exponential fog: attenuates distant objects.')

    # 3.5
    doc.add_heading('3.5 ' + ('Texturizado y environment mapping' if lang == 'es' else 'Texturing and Environment Mapping'), level=2)
    if lang == 'es':
        add_para(doc, 'Se implementó un skybox cubemap con 6 caras PNG de 512×512 píxeles (cielo ártico nublado), renderizado al final del pipeline con la técnica xyww. El shader de agua calcula reflexiones Fresnel dinámicas del cielo con una aproximación elevada a la cuarta potencia.')
    else:
        add_para(doc, 'A cubemap skybox with 6 PNG faces of 512×512 pixels (cloudy Arctic sky) was implemented, rendered at the end of the pipeline using the xyww technique. The water shader computes dynamic Fresnel sky reflections with a fourth-power approximation.')

    # 3.6
    doc.add_heading('3.6 ' + ('Sistema de animaciones' if lang == 'es' else 'Animation System'), level=2)
    anim_h = [('Módulo' if lang == 'es' else 'Module'), ('Nombre' if lang == 'es' else 'Name'),
              ('Tipo' if lang == 'es' else 'Type'), ('Descripción' if lang == 'es' else 'Description')]
    if lang == 'es':
        anim_rows = [
            ['M1_IZQ', 'Deshielo del Ártico', 'Keyframe', 'Iceberg encoge; charco crece'],
            ['M2_IZQ', 'Fauna en Peligro', 'Keyframe', 'Disco de hielo encoge; oso pierde plataforma'],
            ['M3_IZQ', 'Ciudades Inundadas', 'Keyframe', 'Nivel de agua sube cubriendo edificios'],
            ['M1_DER', 'Energía Eólica', 'Keyframe + Continua', 'Turbina acelera; palas rotan'],
            ['M2_DER', 'Auto Eléctrico', 'Keyframe + Continua', 'Color gris→azul; rotación Y'],
            ['M3_DER', 'Captura de Carbono', 'Keyframe + Procedural', 'Árbol crece; humo cíclico'],
            ['M5', 'Acuerdos por el Clima', 'Procedural + Keyframe', 'Globo crece y rota; arcos secuenciales'],
            ['—', 'Olas del océano', 'Procedural', '3 ondas sinusoidales en vertex shader'],
            ['—', 'Caída de nieve', 'Básica', 'Partículas caen con gravedad y deriva'],
            ['—', 'Patrulla de fauna', 'Procedural', 'Pingüinos y lobos con movimiento sinusoidal'],
            ['—', 'Niebla atmosférica', 'Procedural', 'Función exp cuadrática en fragment shader'],
        ]
    else:
        anim_rows = [
            ['M1_IZQ', 'Arctic Melting', 'Keyframe', 'Iceberg shrinks; pool grows'],
            ['M2_IZQ', 'Endangered Fauna', 'Keyframe', 'Ice disc shrinks; bear loses platform'],
            ['M3_IZQ', 'Flooded Cities', 'Keyframe', 'Water level rises covering buildings'],
            ['M1_DER', 'Wind Energy', 'Keyframe + Continuous', 'Turbine accelerates; blades rotate'],
            ['M2_DER', 'Electric Car', 'Keyframe + Continuous', 'Color gray→blue; Y rotation'],
            ['M3_DER', 'Carbon Capture', 'Keyframe + Procedural', 'Tree grows; cyclic smoke'],
            ['M5', 'Climate Agreements', 'Procedural + Keyframe', 'Globe grows and rotates; sequential arcs'],
            ['—', 'Ocean waves', 'Procedural', '3 sinusoidal waves in vertex shader'],
            ['—', 'Snowfall', 'Basic', 'Particles fall with gravity and drift'],
            ['—', 'Fauna patrol', 'Procedural', 'Penguins and wolves with sinusoidal movement'],
            ['—', 'Atmospheric fog', 'Procedural', 'Quadratic exp function in fragment shader'],
        ]
    add_table(doc, anim_h, anim_rows)

    if lang == 'es':
        add_para(doc, 'Tipo 1 — Básica: la caída de nieve aplica una traslación directa en Y cada frame.', bold=True)
        add_para(doc, 'Tipo 2 — Keyframe (LERP): cada módulo usa animT ∈ [0, 1] interpolado en 8 segundos con glm::mix.', bold=True)
        add_para(doc, 'Tipo 3 — Procedural: las olas se generan superponiendo tres sinusoidales en el vertex shader:', bold=True)
    else:
        add_para(doc, 'Type 1 — Basic: snowfall applies a direct Y translation each frame.', bold=True)
        add_para(doc, 'Type 2 — Keyframe (LERP): each module uses animT ∈ [0, 1] interpolated over 8 seconds with glm::mix.', bold=True)
        add_para(doc, 'Type 3 — Procedural: waves are generated by superimposing three sinusoidals in the vertex shader:', bold=True)
    add_code_block(doc, 'float w1 = waveAmp * sin(pos.x * waveFreq + time * waveSpeed);\nfloat w2 = waveAmp * 0.6 * sin(pos.z * waveFreq * 0.7 + time * waveSpeed * 1.3 + 1.2);\nfloat w3 = waveAmp * 0.2 * sin((pos.x+pos.z) * waveFreq * 2.5 + time * waveSpeed * 0.9);\npos.y += w1 + w2 + w3;')

    # 3.7
    doc.add_heading('3.7 ' + ('Diseño narrativo y recorrido' if lang == 'es' else 'Narrative Design and Walkthrough'), level=2)
    if lang == 'es':
        add_para(doc, 'El museo tiene forma de T: vestíbulo (inicio), corredor izquierdo (consecuencias: deshielo, fauna, inundaciones), corredor derecho (soluciones: eólica, auto eléctrico, captura de carbono), y sala final M5 (acuerdos por el clima). Cada módulo se activa con la tecla E dentro de un radio de 5 metros. El progreso se muestra como checklist en la esquina superior derecha.')
    else:
        add_para(doc, 'The museum has a T-shape: lobby (start), left corridor (consequences: melting, fauna, floods), right corridor (solutions: wind, electric car, carbon capture), and final hall M5 (climate agreements). Each module is activated with the E key within a 5-meter radius. Progress is displayed as a checklist in the upper-right corner.')

    # 3.8 Catálogo
    doc.add_heading('3.8 ' + ('Catálogo de assets' if lang == 'es' else 'Asset Catalog'), level=2)

    # Modelos
    doc.add_heading(('Modelos 3D' if lang == 'es' else '3D Models'), level=3)
    model_h = [('Archivo' if lang == 'es' else 'File'), ('Formato' if lang == 'es' else 'Format'),
               ('Módulo' if lang == 'es' else 'Module'), ('Fuente' if lang == 'es' else 'Source'), ('Licencia' if lang == 'es' else 'License')]
    model_rows = [
        ['pine_snow.glb', 'GLB', 'Decoración', 'Quaternius', 'CC0'],
        ['electric_car.glb', 'GLB', 'M2_DER', 'Kenney', 'CC0'],
        ['building_a/b/c.glb', 'GLB', 'M3_IZQ', 'Kenney', 'CC0'],
        ['polar_bear.glb', 'GLB', 'M2_IZQ', 'Quaternius', 'CC0'],
        ['fox.glb', 'GLB', 'Fauna', 'KhronosGroup', 'CC0'],
        ['wolf.glb', 'GLB', 'Fauna', 'Quaternius', 'CC0'],
        ['seal.glb / whale.glb', 'GLB', 'Fauna', 'Quaternius', 'CC0'],
        ['seagull.glb', 'GLB', 'Fauna', 'three.js', 'CC0'],
        ['EolicOBJ.obj', 'OBJ', 'M1_DER', 'Modelo 3D libre', 'CC0'],
        ['Factory.obj', 'OBJ', 'M3_DER', 'Modelo 3D libre', 'CC0'],
        ['igloo_real.obj', 'OBJ', 'Decoración', 'Modelo 3D libre', 'CC0'],
        ['penguin.obj', 'OBJ', 'Fauna', 'Generado', '—'],
    ]
    add_table(doc, model_h, model_rows)

    doc.add_page_break()

    # ── 4. EXPERIMENTOS ──
    doc.add_heading('4. ' + ('Experimentos' if lang == 'es' else 'Experiments'), level=1)
    doc.add_heading('4.1 ' + ('Capturas de los módulos' if lang == 'es' else 'Module Screenshots'), level=2)

    modules = [
        ('VESTIBULO', 'Vestíbulo de entrada' if lang == 'es' else 'Entrance lobby'),
        ('M1_IZQ', 'Deshielo del Ártico' if lang == 'es' else 'Arctic Melting'),
        ('M2_IZQ', 'Fauna en Peligro' if lang == 'es' else 'Endangered Fauna'),
        ('M3_IZQ', 'Ciudades Inundadas' if lang == 'es' else 'Flooded Cities'),
        ('M1_DER', 'Energía Eólica' if lang == 'es' else 'Wind Energy'),
        ('M2_DER', 'Auto Eléctrico' if lang == 'es' else 'Electric Car'),
        ('M3_DER', 'Captura de Carbono' if lang == 'es' else 'Carbon Capture'),
        ('M5', 'Acuerdos por el Clima' if lang == 'es' else 'Climate Agreements'),
    ]
    for i, (mod_id, mod_name) in enumerate(modules):
        fig_n = i + 1
        init_label = 'estado inicial' if lang == 'es' else 'initial state'
        end_label = 'animación completada' if lang == 'es' else 'animation completed'
        add_figure(doc, f'exp_{mod_id}_front.png',
                   f'Figura {fig_n}a. {mod_name} — {init_label}.' if lang == 'es'
                   else f'Figure {fig_n}a. {mod_name} — {init_label}.',
                   width=Inches(5.0))
        if mod_id != 'VESTIBULO':
            add_figure(doc, f'exp_{mod_id}_anim.png',
                       f'Figura {fig_n}b. {mod_name} — {end_label}.' if lang == 'es'
                       else f'Figure {fig_n}b. {mod_name} — {end_label}.',
                       width=Inches(5.0))

    # 4.2 Rendimiento
    doc.add_heading('4.2 ' + ('Rendimiento' if lang == 'es' else 'Performance'), level=2)
    fps_h = [('Zona' if lang == 'es' else 'Zone'), 'FPS ' + ('promedio' if lang == 'es' else 'average'),
             'FPS ' + ('mínimo' if lang == 'es' else 'minimum')]
    fps_rows = [
        [('Vestíbulo' if lang == 'es' else 'Lobby'), '165', '155'],
        [('Corredores' if lang == 'es' else 'Corridors'), '150', '140'],
        [('Sala M5' if lang == 'es' else 'M5 Hall'), '160', '150'],
    ]
    add_table(doc, fps_h, fps_rows)

    # ── 5. CONCLUSIONES ──
    doc.add_heading('5. ' + ('Conclusiones' if lang == 'es' else 'Conclusions'), level=1)
    doc.add_heading(('Conclusión grupal' if lang == 'es' else 'Group Conclusion'), level=2)
    if lang == 'es':
        add_para(doc, 'El desarrollo de este proyecto permitió aplicar de forma integrada los conceptos del curso de Computación Gráfica: desde la configuración del pipeline con OpenGL hasta la implementación de técnicas de iluminación, texturizado, animación y diseño de interacción. La elección de un museo virtual como vehículo narrativo resultó efectiva para dar contexto y propósito a cada técnica gráfica implementada. El tema del calentamiento global no solo le dio sentido al recorrido, sino que nos hizo conscientes de la responsabilidad que tenemos como ingenieros de usar la tecnología para comunicar problemas reales. El proyecto demuestra que la computación gráfica, más allá de ser una herramienta técnica, puede ser un medio de divulgación científica accesible.')
    else:
        add_para(doc, 'The development of this project allowed us to apply the concepts of the Computer Graphics course in an integrated manner: from configuring the pipeline with OpenGL to implementing lighting, texturing, animation and interaction design techniques. The choice of a virtual museum as a narrative vehicle proved effective in providing context and purpose to each implemented graphics technique. The topic of global warming not only gave meaning to the walkthrough but also made us aware of our responsibility as engineers to use technology to communicate real problems. The project demonstrates that computer graphics, beyond being a technical tool, can serve as an accessible medium for scientific communication.')

    doc.add_heading(('Conclusiones individuales' if lang == 'es' else 'Individual Conclusions'), level=2)
    conclusions = {
        'es': [
            ('Navarro Nuño Juan Pablo', 'Desarrollar un museo virtual sobre el calentamiento global me permitió entender cómo la computación gráfica puede ser una herramienta de concientización. Al representar visualmente el deshielo ártico y la pérdida de fauna, el mensaje ambiental se vuelve más tangible que en un texto. Creo que la tecnología gráfica tiene un papel importante en la educación ambiental, y este proyecto me convenció de que vale la pena explorar esa dirección.'),
            ('Martínez Jiménez Israel', 'Mi aportación principal fue la búsqueda e integración de modelos 3D y el sistema de audio. Trabajar con la temática del polo norte me hizo investigar datos que desconocía sobre el impacto real del cambio climático en los ecosistemas árticos. El proyecto me mostró que combinar información científica con experiencias interactivas puede generar una conexión más directa con el usuario que los medios tradicionales.'),
            ('Hernández Cabañas Jesús', 'Implementar las técnicas de iluminación y la arquitectura del sistema me dio una comprensión práctica de los conceptos del curso. Más allá de lo técnico, representar los datos del NSIDC y el IPCC en un entorno 3D interactivo me hizo reflexionar sobre la urgencia del problema climático. Considero que proyectos como este demuestran que la ingeniería puede contribuir a la divulgación de temas ambientales de forma accesible y significativa.'),
        ],
        'en': [
            ('Navarro Nuño Juan Pablo', 'Developing a virtual museum about global warming allowed me to understand how computer graphics can serve as an awareness tool. By visually representing Arctic ice melting and fauna loss, the environmental message becomes more tangible than in a text. I believe graphics technology has an important role in environmental education, and this project convinced me it is worth exploring that direction.'),
            ('Martínez Jiménez Israel', 'My main contribution was the search and integration of 3D models and the audio system. Working with the North Pole theme led me to research data I was unaware of regarding the real impact of climate change on Arctic ecosystems. The project showed me that combining scientific information with interactive experiences can generate a more direct connection with the user than traditional media.'),
            ('Hernández Cabañas Jesús', 'Implementing the lighting techniques and system architecture gave me a practical understanding of the course concepts. Beyond the technical aspects, representing NSIDC and IPCC data in an interactive 3D environment made me reflect on the urgency of the climate problem. I consider that projects like this demonstrate that engineering can contribute to the dissemination of environmental topics in an accessible and meaningful way.'),
        ]
    }
    for name, text in conclusions[lang]:
        p = doc.add_paragraph()
        run = p.add_run(f'{name}: ')
        run.bold = True
        run.font.color.rgb = BLACK
        run.font.name = 'Arial'
        run = p.add_run(text)
        run.font.color.rgb = BLACK
        run.font.name = 'Arial'

    # ── 6. REFERENCIAS ──
    doc.add_heading('6. ' + ('Referencias' if lang == 'es' else 'References'), level=1)
    refs = [
        'Angel, E. y Shreiner, D. (2011). Interactive Computer Graphics: A Top-Down Approach with Shader-Based OpenGL (6.ª ed.). Addison-Wesley.',
        'De Vries, J. (2020). Learn OpenGL: Learn modern OpenGL graphics programming in a step-by-step fashion. https://learnopengl.com/',
        'Hearn, D. y Baker, M. P. (2006). Gráficas por Computadora con OpenGL (3.ª ed.). Pearson Prentice Hall.' if lang == 'es' else 'Hearn, D. & Baker, M. P. (2006). Computer Graphics with OpenGL (3rd ed.). Pearson Prentice Hall.',
        'IEA. (2023). Renewables 2023: Analysis and forecast to 2028. International Energy Agency. https://www.iea.org/reports/renewables-2023',
        'IPCC. (2021). Climate Change 2021: The Physical Science Basis. Cambridge University Press.',
        'IUCN. (2023). Ursus maritimus. The IUCN Red List of Threatened Species. https://www.iucnredlist.org/species/22823/14871490',
        'NSIDC. (2024). Arctic Sea Ice News and Analysis. National Snow and Ice Data Center. https://nsidc.org/arcticseaicenews/',
        'Sylaiou, S., Mania, K., Karoulis, A. y White, M. (2010). Exploring the relationship between presence and enjoyment in a virtual museum. International Journal of Human-Computer Studies, 68(5), 243-253.' if lang == 'es' else 'Sylaiou, S., Mania, K., Karoulis, A. & White, M. (2010). Exploring the relationship between presence and enjoyment in a virtual museum. International Journal of Human-Computer Studies, 68(5), 243-253.',
        'Teodoro-Vite, S. (2026). Materiales del curso Computación Gráfica e Interacción Humano-Computadora. UNAM, Facultad de Ingeniería.' if lang == 'es' else 'Teodoro-Vite, S. (2026). Course materials for Computer Graphics and Human-Computer Interaction. UNAM, Faculty of Engineering.',
        'UNEP. (2023). Emissions Gap Report 2023. United Nations Environment Programme. https://www.unep.org/resources/emissions-gap-report-2023',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(10)
        run.font.color.rgb = BLACK
        run.font.name = 'Arial'
        p.paragraph_format.space_after = Pt(4)


def generate_report(lang='es'):
    doc = Document()
    style_doc(doc)
    build_cover(doc, lang)
    build_bitacora(doc, lang)
    build_body(doc, lang)

    suffix = lang
    docx_path = os.path.join(ENTREGABLES_DIR, f'reporte-{suffix}.docx')
    doc.save(docx_path)
    print(f'DOCX generado: {docx_path}')
    return docx_path


if __name__ == '__main__':
    langs = sys.argv[1] if len(sys.argv) > 1 else 'both'
    if langs in ('es', 'both'):
        generate_report('es')
    if langs in ('en', 'both'):
        generate_report('en')
    print('Listo. Convertir a PDF con LibreOffice.')
