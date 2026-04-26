"""aplicar_cambios.py — Núcleo de edición de la skill /editar-documento.

Toma el último borrador (o la plantilla si N=0), aplica una lista de
cambios estructurada en YAML, y genera v_{N+1}.docx + .diff.md.

Uso:
    python aplicar_cambios.py --caso 2026-001-... --cambios cambios.yml

ESQUEMA del cambios.yml:

    cambios:
      - tipo: sustitucion_texto
        objetivo: "Pedro Pérez"
        nuevo: "Juan García"
        case_sensitive: true
        max_ocurrencias: 0     # 0 = todas
        descripcion: "Cambio de nombre del demandante"

      - tipo: actualizacion_fecha
        objetivo: "{{FECHA_DOCUMENTO}}"
        nueva_fecha: "25 de abril de 2026"
        ambitos: [encabezado, cuerpo, pie]
        descripcion: "Actualización de fecha del documento"

      - tipo: insercion_parrafo
        despues_de: "QUINTA. Causales de rescisión."
        contenido: "SEXTA. Mora superior a 30 días..."
        estilo_de: "QUINTA"     # opcional, copia el style del párrafo de referencia
        descripcion: "Cláusula nueva sobre mora"

      - tipo: eliminacion_parrafo
        contiene: "Texto único del párrafo a eliminar"
        descripcion: "Quitar párrafo obsoleto"

      - tipo: rescritura_seccion
        encabezado: "QUINTA. Causales de rescisión."
        contenido_nuevo: |
          QUINTA. Causales de rescisión.
          Texto nuevo, más estricto, en uno o varios párrafos
          separados por línea en blanco.
        descripcion: "Versión más estricta de la cláusula 5"

      - tipo: formato
        objetivo: "Texto a formatear"
        bold: true
        italic: false
        descripcion: "Negrita en frase clave"

      - tipo: modificar_tabla
        tabla_indice: 0           # 0 = primera tabla del documento
        celda: [1, 1]             # fila 1, columna 1 (0-indexado, 0=encabezado)
        nuevo_texto: "$25,000"
        descripcion: "Renta mensual actualizada"

Salida (stdout, JSON):
    {
      "borrador_generado": "ruta/a/v2-2026-04-25-1635.docx",
      "diff_generado": "ruta/a/v2-2026-04-25-1635.diff.md",
      "version": 2,
      "cambios_aplicados": 4,
      "advertencias": ["..."]
    }
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path

import yaml
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from utils import (
    DESPACHO,
    REPO_ROOT,
    datetime_iso,
    encontrar_caso,
    escribir_caso_yml,
    fecha_iso,
    leer_caso_yml,
    numero_siguiente_version,
    ruta_ultimo_borrador,
    timestamp_borrador,
)


@dataclass
class ResultadoCambio:
    tipo: str
    descripcion: str
    aplicado: bool
    detalle: str
    advertencias: list[str] = field(default_factory=list)


# -------------- Helpers de manipulación DOCX --------------

def _todos_los_parrafos(doc: Document, incluir_headers_footers: bool = True) -> list[Paragraph]:
    """Devuelve todos los párrafos del documento, incluyendo tablas y opcional headers/footers."""
    parrafos: list[Paragraph] = list(doc.paragraphs)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                parrafos.extend(celda.paragraphs)
    if incluir_headers_footers:
        for sec in doc.sections:
            parrafos.extend(sec.header.paragraphs)
            parrafos.extend(sec.footer.paragraphs)
    return parrafos


def _sustituir_en_paragrafo(parrafo: Paragraph, viejo: str, nuevo: str,
                             case_sensitive: bool = True, max_subs: int = 0) -> int:
    """Sustituye texto en un párrafo respetando los runs lo mejor posible.

    Estrategia: si el texto cabe entero en un solo run, modificar ese run
    (preserva formato). Si está partido entre varios runs, se reescribe el
    primer run con el resultado y se borran los demás (pierde formato fino
    en el rango sustituido — aceptable para la mayoría de casos).

    Devuelve número de sustituciones efectuadas.
    """
    flags = 0 if case_sensitive else re.IGNORECASE
    patron = re.compile(re.escape(viejo), flags)

    if not patron.search(parrafo.text):
        return 0

    # Caso simple: texto contenido en un único run
    for run in parrafo.runs:
        if patron.search(run.text):
            count = 0
            def _sub(_m):
                nonlocal count
                if max_subs and count >= max_subs:
                    return _m.group(0)
                count += 1
                return nuevo
            run.text = patron.sub(_sub, run.text)
            if count > 0:
                return count

    # Caso complejo: texto partido entre runs. Reescribir desde texto plano.
    texto_original = parrafo.text
    nuevo_texto, n = patron.subn(nuevo, texto_original, count=max_subs or 0)
    if n == 0:
        return 0
    # Preservar formato del primer run; concatenar todo en él, vaciar los demás
    if parrafo.runs:
        primer_run = parrafo.runs[0]
        primer_run.text = nuevo_texto
        for run in parrafo.runs[1:]:
            run.text = ""
    return n


# -------------- Implementación de los tipos de cambio --------------

def _aplicar_sustitucion_texto(doc: Document, cambio: dict) -> ResultadoCambio:
    objetivo = cambio["objetivo"]
    nuevo = cambio["nuevo"]
    case_sensitive = cambio.get("case_sensitive", True)
    max_oc = cambio.get("max_ocurrencias", 0)
    descripcion = cambio.get("descripcion", f"Sustituir '{objetivo}' por '{nuevo}'")

    total = 0
    restantes = max_oc
    for parrafo in _todos_los_parrafos(doc, incluir_headers_footers=True):
        if max_oc and restantes <= 0:
            break
        n = _sustituir_en_paragrafo(parrafo, objetivo, nuevo, case_sensitive, restantes if max_oc else 0)
        total += n
        if max_oc:
            restantes -= n

    advertencias: list[str] = []
    if total == 0:
        advertencias.append(f"No se encontró '{objetivo}' en el documento")

    return ResultadoCambio(
        tipo="sustitucion_texto",
        descripcion=descripcion,
        aplicado=total > 0,
        detalle=f"{total} ocurrencia(s) sustituida(s)",
        advertencias=advertencias,
    )


def _aplicar_actualizacion_fecha(doc: Document, cambio: dict) -> ResultadoCambio:
    objetivo = cambio["objetivo"]
    nueva = cambio["nueva_fecha"]
    ambitos = cambio.get("ambitos", ["encabezado", "cuerpo", "pie"])
    descripcion = cambio.get("descripcion", f"Actualizar fecha → '{nueva}'")
    flags = re.IGNORECASE if cambio.get("ignorar_mayusculas", False) else 0
    patron = re.compile(re.escape(objetivo), flags)

    total = 0
    if "cuerpo" in ambitos:
        for p in doc.paragraphs:
            total += _sustituir_en_paragrafo(p, objetivo, nueva)
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for p in celda.paragraphs:
                        total += _sustituir_en_paragrafo(p, objetivo, nueva)
    if "encabezado" in ambitos or "pie" in ambitos:
        for sec in doc.sections:
            if "encabezado" in ambitos:
                for p in sec.header.paragraphs:
                    total += _sustituir_en_paragrafo(p, objetivo, nueva)
            if "pie" in ambitos:
                for p in sec.footer.paragraphs:
                    total += _sustituir_en_paragrafo(p, objetivo, nueva)

    advertencias: list[str] = []
    if total == 0:
        advertencias.append(f"No se encontró el patrón de fecha '{objetivo}'")

    return ResultadoCambio(
        tipo="actualizacion_fecha",
        descripcion=descripcion,
        aplicado=total > 0,
        detalle=f"Fecha actualizada en {total} ubicación(es)",
        advertencias=advertencias,
    )


def _aplicar_insercion_parrafo(doc: Document, cambio: dict) -> ResultadoCambio:
    despues_de = cambio["despues_de"]
    contenido = cambio["contenido"]
    descripcion = cambio.get("descripcion", "Inserción de párrafo")
    estilo_de = cambio.get("estilo_de")

    ancla: Paragraph | None = None
    estilo_ancla: Paragraph | None = None
    for p in doc.paragraphs:
        if despues_de in p.text and ancla is None:
            ancla = p
        if estilo_de and estilo_de in p.text and estilo_ancla is None:
            estilo_ancla = p

    if ancla is None:
        return ResultadoCambio(
            tipo="insercion_parrafo",
            descripcion=descripcion,
            aplicado=False,
            detalle="No se encontró el ancla",
            advertencias=[f"Ancla '{despues_de}' no localizada"],
        )

    nuevo_p = deepcopy(ancla._p)
    # Vaciar contenido del clon
    for child in list(nuevo_p):
        if child.tag == qn("w:r"):
            nuevo_p.remove(child)

    ancla._p.addnext(nuevo_p)
    nuevo_parrafo = Paragraph(nuevo_p, ancla._parent)
    nuevo_parrafo.add_run(contenido)
    if estilo_ancla is not None:
        try:
            nuevo_parrafo.style = estilo_ancla.style
        except Exception:
            pass

    return ResultadoCambio(
        tipo="insercion_parrafo",
        descripcion=descripcion,
        aplicado=True,
        detalle=f"Párrafo insertado después de '{despues_de[:40]}...'",
    )


def _aplicar_eliminacion_parrafo(doc: Document, cambio: dict) -> ResultadoCambio:
    contiene = cambio["contiene"]
    descripcion = cambio.get("descripcion", "Eliminación de párrafo")

    eliminados = 0
    for p in list(doc.paragraphs):
        if contiene in p.text:
            p._element.getparent().remove(p._element)
            eliminados += 1

    advertencias: list[str] = []
    if eliminados == 0:
        advertencias.append(f"Ningún párrafo contiene '{contiene}'")

    return ResultadoCambio(
        tipo="eliminacion_parrafo",
        descripcion=descripcion,
        aplicado=eliminados > 0,
        detalle=f"{eliminados} párrafo(s) eliminado(s)",
        advertencias=advertencias,
    )


def _aplicar_rescritura_seccion(doc: Document, cambio: dict) -> ResultadoCambio:
    encabezado = cambio["encabezado"]
    contenido_nuevo = cambio["contenido_nuevo"]
    descripcion = cambio.get("descripcion", f"Reescritura de '{encabezado}'")

    inicio_idx: int | None = None
    parrafos = list(doc.paragraphs)
    for i, p in enumerate(parrafos):
        if encabezado in p.text:
            inicio_idx = i
            break
    if inicio_idx is None:
        return ResultadoCambio(
            tipo="rescritura_seccion",
            descripcion=descripcion,
            aplicado=False,
            detalle="Encabezado no localizado",
            advertencias=[f"No se encontró '{encabezado}'"],
        )

    # Detectar fin de la sección: siguiente párrafo que comience con un patrón de cláusula
    patron_clausula = re.compile(r"^(PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|SÉPTIMA|OCTAVA|NOVENA|DÉCIMA|UNDÉCIMA|DUODÉCIMA|\d+\.)", re.IGNORECASE)
    fin_idx = len(parrafos)
    for i in range(inicio_idx + 1, len(parrafos)):
        if patron_clausula.match(parrafos[i].text.strip()):
            fin_idx = i
            break

    # Eliminar párrafos del inicio (incluido) hasta fin_idx (exclusivo)
    estilo_ancla = parrafos[inicio_idx].style
    for p in parrafos[inicio_idx:fin_idx]:
        p._element.getparent().remove(p._element)

    # Insertar el contenido nuevo, respetando saltos en líneas en blanco
    if fin_idx < len(parrafos):
        ancla_post = parrafos[fin_idx]
        bloques = [b for b in contenido_nuevo.strip().split("\n\n") if b.strip()]
        for bloque in bloques:
            nuevo_p = deepcopy(ancla_post._p)
            for child in list(nuevo_p):
                if child.tag == qn("w:r"):
                    nuevo_p.remove(child)
            ancla_post._p.addprevious(nuevo_p)
            np = Paragraph(nuevo_p, ancla_post._parent)
            np.add_run(bloque.strip())
            try:
                np.style = estilo_ancla
            except Exception:
                pass
    else:
        # Si no hay ancla posterior, agregar al final del body
        for bloque in contenido_nuevo.strip().split("\n\n"):
            if bloque.strip():
                p = doc.add_paragraph(bloque.strip())
                try:
                    p.style = estilo_ancla
                except Exception:
                    pass

    return ResultadoCambio(
        tipo="rescritura_seccion",
        descripcion=descripcion,
        aplicado=True,
        detalle=f"Sección '{encabezado[:40]}...' reescrita ({fin_idx - inicio_idx} párrafos reemplazados)",
    )


def _aplicar_formato(doc: Document, cambio: dict) -> ResultadoCambio:
    objetivo = cambio["objetivo"]
    descripcion = cambio.get("descripcion", "Cambio de formato")
    bold = cambio.get("bold")
    italic = cambio.get("italic")
    underline = cambio.get("underline")

    aplicados = 0
    for parrafo in _todos_los_parrafos(doc, incluir_headers_footers=False):
        if objetivo not in parrafo.text:
            continue
        for run in parrafo.runs:
            if objetivo in run.text:
                if bold is not None:
                    run.bold = bold
                if italic is not None:
                    run.italic = italic
                if underline is not None:
                    run.underline = underline
                aplicados += 1

    advertencias: list[str] = []
    if aplicados == 0:
        advertencias.append(f"No se encontró '{objetivo}' en runs individuales")

    return ResultadoCambio(
        tipo="formato",
        descripcion=descripcion,
        aplicado=aplicados > 0,
        detalle=f"Formato aplicado a {aplicados} run(s)",
        advertencias=advertencias,
    )


def _aplicar_modificar_tabla(doc: Document, cambio: dict) -> ResultadoCambio:
    indice_tabla = cambio.get("tabla_indice", 0)
    celda = cambio["celda"]
    nuevo_texto = cambio["nuevo_texto"]
    descripcion = cambio.get("descripcion", f"Modificar celda {celda} de tabla {indice_tabla}")

    if indice_tabla >= len(doc.tables):
        return ResultadoCambio(
            tipo="modificar_tabla",
            descripcion=descripcion,
            aplicado=False,
            detalle=f"Tabla {indice_tabla} no existe (hay {len(doc.tables)})",
            advertencias=["Índice de tabla fuera de rango"],
        )

    tabla = doc.tables[indice_tabla]
    fila, col = celda
    if fila >= len(tabla.rows) or col >= len(tabla.rows[fila].cells):
        return ResultadoCambio(
            tipo="modificar_tabla",
            descripcion=descripcion,
            aplicado=False,
            detalle="Celda fuera de rango",
            advertencias=[f"Celda [{fila},{col}] no existe"],
        )

    cell = tabla.rows[fila].cells[col]
    # Preservar formato: editar primer run del primer párrafo
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = nuevo_texto
        for run in cell.paragraphs[0].runs[1:]:
            run.text = ""
    else:
        cell.text = nuevo_texto

    return ResultadoCambio(
        tipo="modificar_tabla",
        descripcion=descripcion,
        aplicado=True,
        detalle=f"Celda [{fila},{col}] de tabla {indice_tabla} actualizada",
    )


def _aplicar_agregar_tabla(doc: Document, cambio: dict) -> ResultadoCambio:
    descripcion = cambio.get("descripcion", "Agregar tabla nueva")
    filas = cambio["filas"]  # lista de listas (matriz)
    despues_de = cambio.get("despues_de")  # opcional: ancla por texto

    if not filas:
        return ResultadoCambio(
            tipo="agregar_tabla",
            descripcion=descripcion,
            aplicado=False,
            detalle="Lista de filas vacía",
            advertencias=["Se requiere al menos una fila"],
        )

    n_cols = max(len(f) for f in filas)
    nueva = doc.add_table(rows=len(filas), cols=n_cols)
    nueva.style = "Table Grid"
    for i, fila in enumerate(filas):
        for j, valor in enumerate(fila):
            nueva.cell(i, j).text = str(valor)

    # Si hay ancla, mover la tabla recién agregada
    if despues_de:
        for p in doc.paragraphs:
            if despues_de in p.text:
                p._p.addnext(nueva._element)
                break

    return ResultadoCambio(
        tipo="agregar_tabla",
        descripcion=descripcion,
        aplicado=True,
        detalle=f"Tabla agregada ({len(filas)} filas × {n_cols} columnas)",
    )


_DISPATCHER = {
    "sustitucion_texto": _aplicar_sustitucion_texto,
    "actualizacion_fecha": _aplicar_actualizacion_fecha,
    "insercion_parrafo": _aplicar_insercion_parrafo,
    "eliminacion_parrafo": _aplicar_eliminacion_parrafo,
    "rescritura_seccion": _aplicar_rescritura_seccion,
    "formato": _aplicar_formato,
    "modificar_tabla": _aplicar_modificar_tabla,
    "agregar_tabla": _aplicar_agregar_tabla,
}


# -------------- Generador de .diff.md --------------

def generar_diff_md(version: int, fuente: str, resultados: list[ResultadoCambio],
                    obs_ortografia: list[dict] | None = None) -> str:
    lineas: list[str] = []
    lineas.append(f"# Diff — borrador v{version} ({datetime_iso()})")
    lineas.append("")
    if version == 1:
        lineas.append(f"> Cambios aplicados respecto a la **plantilla de origen** (`{fuente}`).")
    else:
        lineas.append(f"> Cambios aplicados respecto al **borrador v{version - 1}** (`{fuente}`).")
    lineas.append("")

    aplicados = [r for r in resultados if r.aplicado]
    fallidos = [r for r in resultados if not r.aplicado]

    lineas.append("## Resumen ejecutivo")
    lineas.append("")
    lineas.append(f"- Cambios solicitados: **{len(resultados)}**")
    lineas.append(f"- Cambios aplicados:   **{len(aplicados)}**")
    if fallidos:
        lineas.append(f"- Cambios fallidos:    **{len(fallidos)}** (ver advertencias)")
    lineas.append("")

    lineas.append("## Cambios aplicados")
    lineas.append("")
    if not aplicados:
        lineas.append("(ninguno)")
    else:
        lineas.append("| # | Tipo | Descripción | Detalle |")
        lineas.append("|---|------|-------------|---------|")
        for i, r in enumerate(aplicados, 1):
            lineas.append(f"| {i} | `{r.tipo}` | {r.descripcion} | {r.detalle} |")
    lineas.append("")

    if fallidos:
        lineas.append("## Cambios no aplicados")
        lineas.append("")
        for r in fallidos:
            lineas.append(f"- **{r.descripcion}** (`{r.tipo}`): {r.detalle}")
            for adv in r.advertencias:
                lineas.append(f"  - ⚠️ {adv}")
        lineas.append("")

    advertencias_globales = [adv for r in resultados for adv in r.advertencias]
    if advertencias_globales and not fallidos:
        lineas.append("## Advertencias")
        lineas.append("")
        for adv in advertencias_globales:
            lineas.append(f"- ⚠️ {adv}")
        lineas.append("")

    if obs_ortografia is not None:
        lineas.append("## Validación ortográfica")
        lineas.append("")
        if not obs_ortografia:
            lineas.append("Sin observaciones.")
        else:
            lineas.append(f"Se detectaron **{len(obs_ortografia)}** observaciones (no bloqueantes):")
            lineas.append("")
            for o in obs_ortografia[:20]:
                lineas.append(f"- `{o.get('texto', '')}` → sugerencias: {o.get('sugerencias', [])}")
            if len(obs_ortografia) > 20:
                lineas.append(f"- ... y {len(obs_ortografia) - 20} más")
        lineas.append("")

    return "\n".join(lineas)


# -------------- Main --------------

def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--caso", required=True, help="ID del caso (carpeta en despacho/casos/)")
    parser.add_argument("--cambios", required=True, help="Ruta al archivo YAML con la lista de cambios")
    args = parser.parse_args()

    caso_dir = encontrar_caso(args.caso)
    if caso_dir is None:
        sys.stderr.write(f"ERROR: caso '{args.caso}' no encontrado\n")
        return 1

    cambios_path = Path(args.cambios)
    if not cambios_path.exists():
        sys.stderr.write(f"ERROR: archivo de cambios '{cambios_path}' no encontrado\n")
        return 1

    with cambios_path.open(encoding="utf-8") as f:
        spec = yaml.safe_load(f)
    cambios = spec.get("cambios", [])
    if not cambios:
        sys.stderr.write("ERROR: el YAML no contiene 'cambios' o está vacío\n")
        return 1

    # Determinar fuente: último borrador o plantilla
    ultimo = ruta_ultimo_borrador(caso_dir)
    if ultimo is None:
        fuente = caso_dir / "01-plantilla" / "plantilla-origen.docx"
        if not fuente.exists():
            sys.stderr.write(f"ERROR: no hay plantilla en {fuente}\n")
            return 1
    else:
        fuente = ultimo

    nueva_version = numero_siguiente_version(caso_dir)
    ts = timestamp_borrador()
    salida_docx = caso_dir / "02-borradores" / f"v{nueva_version}-{ts}.docx"
    salida_diff = caso_dir / "02-borradores" / f"v{nueva_version}-{ts}.diff.md"

    # Aplicar cambios
    doc = Document(str(fuente))
    resultados: list[ResultadoCambio] = []
    for cambio in cambios:
        tipo = cambio.get("tipo")
        if tipo not in _DISPATCHER:
            resultados.append(ResultadoCambio(
                tipo=tipo or "desconocido",
                descripcion=cambio.get("descripcion", "(sin descripción)"),
                aplicado=False,
                detalle=f"Tipo '{tipo}' no soportado",
                advertencias=[f"Tipos válidos: {sorted(_DISPATCHER.keys())}"],
            ))
            continue
        try:
            r = _DISPATCHER[tipo](doc, cambio)
        except Exception as e:
            r = ResultadoCambio(
                tipo=tipo,
                descripcion=cambio.get("descripcion", "(sin descripción)"),
                aplicado=False,
                detalle=f"Excepción: {type(e).__name__}: {e}",
                advertencias=[],
            )
        resultados.append(r)

    doc.save(str(salida_docx))

    # Generar diff
    diff_md = generar_diff_md(
        version=nueva_version,
        fuente=fuente.name,
        resultados=resultados,
    )
    salida_diff.write_text(diff_md, encoding="utf-8")

    # Actualizar caso.yml
    datos_caso = leer_caso_yml(caso_dir)
    datos_caso.setdefault("borradores", []).append({
        "version": nueva_version,
        "archivo": salida_docx.name,
        "fecha": datetime_iso(),
        "resumen": ", ".join(r.descripcion for r in resultados if r.aplicado)[:200] or "(sin cambios aplicados)",
    })
    datos_caso["fechas"]["ultima_modificacion"] = fecha_iso()
    escribir_caso_yml(caso_dir, datos_caso)

    advertencias_totales = [adv for r in resultados for adv in r.advertencias]
    resultado = {
        "borrador_generado": str(salida_docx.relative_to(REPO_ROOT)).replace("\\", "/"),
        "diff_generado": str(salida_diff.relative_to(REPO_ROOT)).replace("\\", "/"),
        "version": nueva_version,
        "cambios_solicitados": len(resultados),
        "cambios_aplicados": sum(1 for r in resultados if r.aplicado),
        "advertencias": advertencias_totales,
    }
    print(json.dumps(resultado, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
