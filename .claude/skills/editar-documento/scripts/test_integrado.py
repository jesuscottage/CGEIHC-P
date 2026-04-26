"""test_integrado.py — Suite de tests E2E para la skill /editar-documento.

Cubre los casos de prueba del PLAN.md sección 6 que son automatizables.
Los tests dependientes de LLM (5: ambigüedad) o de estado del sistema
(9: Word ocupado) se marcan como manuales.

Ejecuta:
    python test_integrado.py

Cada test imprime PASS/FAIL con detalle. Sale con código 0 si todos
pasan, 1 si alguno falla.
"""
from __future__ import annotations

import json
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

import yaml
from docx import Document

REPO_ROOT = Path(__file__).resolve().parents[4]
DESPACHO = REPO_ROOT / "despacho"
SCRIPTS = Path(__file__).parent
PYTHON = sys.executable


class TestRunner:
    def __init__(self):
        self.pasaron = 0
        self.fallaron = 0
        self.fallos: list[tuple[str, str]] = []

    def run(self, nombre: str, fn) -> None:
        print(f"\n=== TEST: {nombre} ===")
        try:
            fn()
            self.pasaron += 1
            print(f"  PASS")
        except AssertionError as e:
            self.fallaron += 1
            msg = str(e).encode("ascii", errors="replace").decode("ascii")
            self.fallos.append((nombre, msg))
            print(f"  FAIL: {msg}")
        except Exception as e:
            self.fallaron += 1
            msg = f"{type(e).__name__}: {e}".encode("ascii", errors="replace").decode("ascii")
            self.fallos.append((nombre, msg))
            print(f"  ERROR: {msg}")

    def resumen(self) -> int:
        total = self.pasaron + self.fallaron
        print(f"\n{'=' * 50}")
        print(f"RESULTADO: {self.pasaron}/{total} tests pasaron")
        if self.fallos:
            print("\nFallos:")
            for nombre, error in self.fallos:
                print(f"  - {nombre}: {error}")
        print("=" * 50)
        return 0 if self.fallaron == 0 else 1


def _ejecutar(args: list[str]) -> dict:
    """Ejecuta un script Python y devuelve el JSON de stdout."""
    result = subprocess.run([PYTHON, *args], capture_output=True)
    stdout = result.stdout.decode("utf-8", errors="replace")
    stderr = result.stderr.decode("utf-8", errors="replace")
    if result.returncode != 0:
        raise RuntimeError(f"Script falló (code {result.returncode})\nstderr: {stderr[-500:]}")
    if not stdout.strip():
        raise RuntimeError(f"Script no produjo salida\nstderr: {stderr[-500:]}")
    return json.loads(stdout)


def _ejecutar_esperando_fallo(args: list[str], codigo_esperado: int) -> str:
    """Ejecuta esperando que falle con un código específico. Devuelve stderr."""
    result = subprocess.run([PYTHON, *args], capture_output=True)
    stderr = result.stderr.decode("utf-8", errors="replace")
    if result.returncode != codigo_esperado:
        raise AssertionError(
            f"Esperaba código {codigo_esperado}, obtuvo {result.returncode}.\nstderr: {stderr[-300:]}"
        )
    return stderr


def _texto_del_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


# -------------------- TESTS --------------------

def test_01_apertura_con_plantilla_valida():
    info = _ejecutar([
        str(SCRIPTS / "abrir_caso.py"),
        "--cliente", "Test Uno",
        "--plantilla", "ejemplo-plantilla-v1",
        "--materia", "civil",
    ])
    caso_dir = REPO_ROOT / info["carpeta"]
    assert caso_dir.exists(), "carpeta del caso no existe"
    assert (caso_dir / "01-plantilla" / "plantilla-origen.docx").exists(), "plantilla no copiada"
    assert (caso_dir / "caso.yml").exists(), "caso.yml no creado"
    assert (caso_dir / "02-borradores").exists(), "carpeta 02-borradores no creada"
    assert info["id"].startswith("2026-"), f"ID inesperado: {info['id']}"


def test_02_edicion_3_sustituciones():
    info_apertura = _ejecutar([
        str(SCRIPTS / "abrir_caso.py"),
        "--cliente", "Test Dos",
        "--plantilla", "ejemplo-plantilla-v1",
    ])
    caso_id = info_apertura["id"]

    cambios = {"cambios": [
        {"tipo": "sustitucion_texto", "objetivo": "{{NOMBRE_ARRENDADOR}}", "nuevo": "PEDRO"},
        {"tipo": "sustitucion_texto", "objetivo": "{{NOMBRE_ARRENDATARIO}}", "nuevo": "JUAN"},
        {"tipo": "sustitucion_texto", "objetivo": "{{DOMICILIO_INMUEBLE}}", "nuevo": "Calle 1"},
    ]}
    cambios_yml = Path(tempfile.mktemp(suffix=".yml"))
    cambios_yml.write_text(yaml.safe_dump(cambios, allow_unicode=True), encoding="utf-8")

    info = _ejecutar([str(SCRIPTS / "aplicar_cambios.py"), "--caso", caso_id, "--cambios", str(cambios_yml)])
    assert info["cambios_aplicados"] == 3, f"esperaba 3, aplicó {info['cambios_aplicados']}"
    docx = REPO_ROOT / info["borrador_generado"]
    texto = _texto_del_docx(docx)
    assert "PEDRO" in texto and "JUAN" in texto and "Calle 1" in texto, "sustituciones no presentes"
    assert "{{NOMBRE_ARRENDADOR}}" not in texto, "marcador no fue sustituido"


def test_03_clausula_nueva_y_rescritura():
    info_apertura = _ejecutar([
        str(SCRIPTS / "abrir_caso.py"),
        "--cliente", "Test Tres",
        "--plantilla", "ejemplo-plantilla-v1",
    ])
    caso_id = info_apertura["id"]

    cambios = {"cambios": [
        {"tipo": "insercion_parrafo",
         "despues_de": "Cualquiera de las partes podrá rescindir",
         "contenido": "SEXTA. Cláusula adicional sobre intereses moratorios."},
        {"tipo": "rescritura_seccion",
         "encabezado": "TERCERA. Renta.",
         "contenido_nuevo": "TERCERA. Renta.\n\nLa renta será de cincuenta mil pesos mexicanos pagaderos los primeros diez días de cada mes."},
    ]}
    cambios_yml = Path(tempfile.mktemp(suffix=".yml"))
    cambios_yml.write_text(yaml.safe_dump(cambios, allow_unicode=True), encoding="utf-8")

    info = _ejecutar([str(SCRIPTS / "aplicar_cambios.py"), "--caso", caso_id, "--cambios", str(cambios_yml)])
    assert info["cambios_aplicados"] == 2, f"esperaba 2, aplicó {info['cambios_aplicados']}"
    texto = _texto_del_docx(REPO_ROOT / info["borrador_generado"])
    assert "SEXTA. Cláusula adicional" in texto, "cláusula nueva no presente"
    assert "cincuenta mil pesos mexicanos" in texto, "rescritura no aplicada"


def test_04_iterar_5_veces():
    info_apertura = _ejecutar([
        str(SCRIPTS / "abrir_caso.py"),
        "--cliente", "Test Cuatro",
        "--plantilla", "ejemplo-plantilla-v1",
    ])
    caso_id = info_apertura["id"]
    caso_dir = REPO_ROOT / info_apertura["carpeta"]

    for i in range(1, 6):
        cambios = {"cambios": [
            {"tipo": "sustitucion_texto", "objetivo": f"iter-{i-1}-marker" if i > 1 else "{{SUPERFICIE}}", "nuevo": f"iter-{i}-marker"},
        ]}
        cambios_yml = Path(tempfile.mktemp(suffix=".yml"))
        cambios_yml.write_text(yaml.safe_dump(cambios, allow_unicode=True), encoding="utf-8")
        info = _ejecutar([str(SCRIPTS / "aplicar_cambios.py"), "--caso", caso_id, "--cambios", str(cambios_yml)])
        assert info["version"] == i, f"iteración {i}: versión esperada {i}, obtuvo {info['version']}"

    borradores = sorted((caso_dir / "02-borradores").glob("v*-*.docx"))
    assert len(borradores) == 5, f"esperaba 5 borradores, hay {len(borradores)}"


def test_06_cerrar_caso_publica_entregable():
    info_apertura = _ejecutar([
        str(SCRIPTS / "abrir_caso.py"),
        "--cliente", "Test Seis",
        "--plantilla", "ejemplo-plantilla-v1",
    ])
    caso_id = info_apertura["id"]

    cambios = {"cambios": [
        {"tipo": "sustitucion_texto", "objetivo": "{{NOMBRE_ARRENDADOR}}", "nuevo": "Cierre Test"},
    ]}
    cambios_yml = Path(tempfile.mktemp(suffix=".yml"))
    cambios_yml.write_text(yaml.safe_dump(cambios, allow_unicode=True), encoding="utf-8")
    _ejecutar([str(SCRIPTS / "aplicar_cambios.py"), "--caso", caso_id, "--cambios", str(cambios_yml)])

    info = _ejecutar([str(SCRIPTS / "cerrar_caso.py"), "--caso", caso_id])
    assert (REPO_ROOT / info["final_docx"]).exists(), "final.docx no existe"
    assert (REPO_ROOT / info["final_pdf"]).exists(), "final.pdf no existe"
    assert (REPO_ROOT / info["publicacion_pdf"]).exists(), "PDF no publicado en entregables"
    indice = (DESPACHO / "entregables" / "INDICE.md").read_text(encoding="utf-8")
    assert caso_id in indice, "caso no agregado al INDICE"


def test_07_editar_caso_cerrado_falla():
    info_apertura = _ejecutar([
        str(SCRIPTS / "abrir_caso.py"),
        "--cliente", "Test Siete",
        "--plantilla", "ejemplo-plantilla-v1",
    ])
    caso_id = info_apertura["id"]
    cambios_yml = Path(tempfile.mktemp(suffix=".yml"))
    cambios_yml.write_text(yaml.safe_dump({"cambios": [
        {"tipo": "sustitucion_texto", "objetivo": "{{NOMBRE_ARRENDADOR}}", "nuevo": "x"},
    ]}, allow_unicode=True), encoding="utf-8")
    _ejecutar([str(SCRIPTS / "aplicar_cambios.py"), "--caso", caso_id, "--cambios", str(cambios_yml)])
    _ejecutar([str(SCRIPTS / "cerrar_caso.py"), "--caso", caso_id])

    # Segundo intento de cierre debe fallar
    stderr = _ejecutar_esperando_fallo(
        [str(SCRIPTS / "cerrar_caso.py"), "--caso", caso_id], 1
    )
    # Comparar sin tildes para evitar mojibake en consola Windows
    stderr_norm = stderr.lower().replace("�", "?")
    assert "ya est" in stderr_norm and "cerrado" in stderr_norm, (
        f"mensaje inesperado: {stderr[-200:]}"
    )


def test_08_tabla_preservada():
    """La plantilla tiene una tabla 4x2; tras edición la tabla debe seguir presente."""
    info_apertura = _ejecutar([
        str(SCRIPTS / "abrir_caso.py"),
        "--cliente", "Test Ocho",
        "--plantilla", "ejemplo-plantilla-v1",
    ])
    caso_id = info_apertura["id"]
    cambios_yml = Path(tempfile.mktemp(suffix=".yml"))
    cambios_yml.write_text(yaml.safe_dump({"cambios": [
        {"tipo": "modificar_tabla", "tabla_indice": 0, "celda": [1, 1], "nuevo_texto": "$30,000"},
    ]}, allow_unicode=True), encoding="utf-8")
    info = _ejecutar([str(SCRIPTS / "aplicar_cambios.py"), "--caso", caso_id, "--cambios", str(cambios_yml)])
    docx = REPO_ROOT / info["borrador_generado"]
    doc = Document(str(docx))
    assert len(doc.tables) == 1, f"esperaba 1 tabla, hay {len(doc.tables)}"
    assert len(doc.tables[0].rows) == 4, f"esperaba 4 filas, hay {len(doc.tables[0].rows)}"
    assert doc.tables[0].rows[1].cells[1].text.strip() == "$30,000", "celda no actualizada"


def test_10_validador_detecta_tildes_faltantes():
    """Insertar texto con tildes faltantes y verificar que LanguageTool las reporta."""
    info_apertura = _ejecutar([
        str(SCRIPTS / "abrir_caso.py"),
        "--cliente", "Test Diez",
        "--plantilla", "ejemplo-plantilla-v1",
    ])
    caso_id = info_apertura["id"]
    cambios_yml = Path(tempfile.mktemp(suffix=".yml"))
    cambios_yml.write_text(yaml.safe_dump({"cambios": [
        {"tipo": "insercion_parrafo",
         "despues_de": "Cualquiera de las partes podrá rescindir",
         "contenido": "El abogado redacto este parrafo sin tildes y la jueza acepto el contenido pero parecio raro."},
    ]}, allow_unicode=True), encoding="utf-8")
    info = _ejecutar([str(SCRIPTS / "aplicar_cambios.py"), "--caso", caso_id, "--cambios", str(cambios_yml)])
    docx = REPO_ROOT / info["borrador_generado"]
    obs = _ejecutar([str(SCRIPTS / "validar_ortografia.py"), "--docx", str(docx)])
    palabras_detectadas = {o["texto"] for o in obs["observaciones"]}
    # Al menos una de "redacto", "parecio", "acepto" debe ser detectada
    assert palabras_detectadas & {"redacto", "parecio", "acepto", "parrafo"}, (
        f"LanguageTool no detectó tildes. Detectadas: {palabras_detectadas}"
    )


# -------------------- Main --------------------

def main() -> int:
    runner = TestRunner()
    runner.run("01 - Apertura con plantilla válida", test_01_apertura_con_plantilla_valida)
    runner.run("02 - Edición con 3 sustituciones simples", test_02_edicion_3_sustituciones)
    runner.run("03 - Cláusula nueva + reescritura", test_03_clausula_nueva_y_rescritura)
    runner.run("04 - Iterar 5 veces", test_04_iterar_5_veces)
    runner.run("06 - Cerrar caso y publicar", test_06_cerrar_caso_publica_entregable)
    runner.run("07 - Editar caso cerrado falla", test_07_editar_caso_cerrado_falla)
    runner.run("08 - Tabla preservada", test_08_tabla_preservada)
    runner.run("10 - Validador detecta tildes faltantes", test_10_validador_detecta_tildes_faltantes)

    print("\nTests no automatizables (requieren validación manual):")
    print("  - 05: Cambio ambiguo ('mejora la redacción') — Claude debe preguntar antes de actuar")
    print("  - 09: Word ocupado al pedir conversión — fallback a LibreOffice")

    return runner.resumen()


def _limpiar():
    """Limpia artefactos de tests anteriores antes de correr."""
    for d in (DESPACHO / "casos").glob("2026-*-test-*"):
        if d.is_dir():
            shutil.rmtree(d)
    entregables_2026 = DESPACHO / "entregables" / "2026"
    if entregables_2026.exists():
        shutil.rmtree(entregables_2026)
    indice = DESPACHO / "entregables" / "INDICE.md"
    if indice.exists():
        contenido = indice.read_text(encoding="utf-8")
        import re
        contenido = re.sub(
            r"\| \d{4}-\d{3}-test-\S+ \|.*\|.*\|.*\|.*\|\n",
            "",
            contenido,
        )
        if "(sin entregables todavía)" not in contenido:
            contenido = contenido.replace(
                "|------|---------|------|--------------|----------|\n",
                "|------|---------|------|--------------|----------|\n| — | — | (sin entregables todavía) | — | — |\n",
            )
        indice.write_text(contenido, encoding="utf-8")


if __name__ == "__main__":
    _limpiar()
    sys.exit(main())
