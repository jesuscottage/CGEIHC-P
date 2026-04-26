"""utils.py — helpers compartidos por los scripts de la skill /editar-documento.

Funciones:
- slugify(texto): produce un slug kebab-case sin acentos, máx. 30 chars
- siguiente_correlativo(año, casos_dir): asigna el siguiente NNN del año
- leer_caso_yml(caso_dir): carga caso.yml como dict
- escribir_caso_yml(caso_dir, datos): escribe caso.yml preservando legibilidad
- ruta_ultimo_borrador(caso_dir): devuelve Path del v_N más reciente
- timestamp_borrador(): "2026-04-25-1635" para nombres de archivo
- extraer_texto(docx_path): extrae texto plano del DOCX (para que Claude lo lea)

CLI: este módulo expone subcomandos para invocación directa desde Bash.
    python utils.py slug "Juan García López"
    python utils.py extraer-texto archivo.docx
"""
from __future__ import annotations

import argparse
import json
import re
import sys
import unicodedata
from datetime import datetime
from pathlib import Path

import yaml
from docx import Document

REPO_ROOT = Path(__file__).resolve().parents[4]
DESPACHO = REPO_ROOT / "despacho"


def slugify(texto: str, max_len: int = 30) -> str:
    """Convierte 'Juan García López' → 'juan-garcia-lopez'.

    Quita acentos, baja a minúsculas, deja sólo [a-z0-9-], colapsa
    guiones repetidos y trunca a max_len sin romper palabras.
    """
    nfkd = unicodedata.normalize("NFKD", texto)
    sin_acentos = "".join(c for c in nfkd if not unicodedata.combining(c))
    en_minusculas = sin_acentos.lower()
    solo_alnum = re.sub(r"[^a-z0-9]+", "-", en_minusculas)
    limpio = solo_alnum.strip("-")

    if len(limpio) <= max_len:
        return limpio
    truncado = limpio[:max_len].rsplit("-", 1)[0]
    return truncado or limpio[:max_len]


def siguiente_correlativo(año: int, casos_dir: Path) -> int:
    """Devuelve el siguiente NNN para el año dado (1 si es el primero).

    Excluye el caso de ejemplo (000) del cálculo.
    """
    if not casos_dir.exists():
        return 1
    prefijo = f"{año}-"
    correlativos = []
    for d in casos_dir.iterdir():
        if not d.is_dir() or not d.name.startswith(prefijo):
            continue
        partes = d.name.split("-", 2)
        if len(partes) < 2:
            continue
        try:
            num = int(partes[1])
            if num > 0:
                correlativos.append(num)
        except ValueError:
            continue
    return (max(correlativos) + 1) if correlativos else 1


def leer_caso_yml(caso_dir: Path) -> dict:
    archivo = caso_dir / "caso.yml"
    if not archivo.exists():
        raise FileNotFoundError(f"No existe {archivo}")
    with archivo.open(encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


def escribir_caso_yml(caso_dir: Path, datos: dict) -> None:
    archivo = caso_dir / "caso.yml"
    with archivo.open("w", encoding="utf-8") as f:
        yaml.safe_dump(datos, f, allow_unicode=True, sort_keys=False, indent=2, default_flow_style=False)


def ruta_ultimo_borrador(caso_dir: Path) -> Path | None:
    """Devuelve la ruta del borrador con N más alto, o None si no hay ninguno."""
    borradores_dir = caso_dir / "02-borradores"
    if not borradores_dir.exists():
        return None
    docs = sorted(borradores_dir.glob("v*-*.docx"))
    if not docs:
        return None
    def numero_version(p: Path) -> int:
        m = re.match(r"v(\d+)-", p.name)
        return int(m.group(1)) if m else 0
    return max(docs, key=numero_version)


def numero_siguiente_version(caso_dir: Path) -> int:
    actual = ruta_ultimo_borrador(caso_dir)
    if actual is None:
        return 1
    m = re.match(r"v(\d+)-", actual.name)
    return (int(m.group(1)) + 1) if m else 1


def timestamp_borrador() -> str:
    return datetime.now().strftime("%Y-%m-%d-%H%M")


def fecha_iso() -> str:
    return datetime.now().strftime("%Y-%m-%d")


def datetime_iso() -> str:
    return datetime.now().strftime("%Y-%m-%dT%H:%M:%S")


def extraer_texto(docx_path: Path) -> str:
    """Extrae texto plano del DOCX, preservando estructura básica.

    Útil para que Claude lea el contenido del último borrador antes de
    decidir qué cambios aplicar. Itera sólo sobre los hijos directos del
    body para evitar duplicar el contenido de las tablas.
    """
    from docx.text.paragraph import Paragraph
    doc = Document(str(docx_path))
    bloques: list[str] = []

    for sec in doc.sections:
        for p in sec.header.paragraphs:
            if p.text.strip():
                bloques.append(f"[ENCABEZADO] {p.text}")
        for p in sec.footer.paragraphs:
            if p.text.strip():
                bloques.append(f"[PIE] {p.text}")

    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for elem in doc.element.body:
        tag = elem.tag
        if tag == f"{W}p":
            parrafo = Paragraph(elem, doc)
            if parrafo.text.strip():
                bloques.append(parrafo.text)
        elif tag == f"{W}tbl":
            bloques.append("[TABLA]")
            for row in elem.iter(f"{W}tr"):
                celdas = []
                for tc in row.iter(f"{W}tc"):
                    texto_celda = "".join(t.text or "" for t in tc.iter(f"{W}t"))
                    celdas.append(texto_celda)
                bloques.append(" | ".join(celdas))
            bloques.append("[/TABLA]")
    return "\n".join(bloques)


def listar_plantillas() -> list[dict]:
    """Lista plantillas disponibles leyendo nombre + meta.yml hermano."""
    plantillas_dir = DESPACHO / "plantillas"
    resultado: list[dict] = []
    if not plantillas_dir.exists():
        return resultado
    for archivo in plantillas_dir.rglob("*.docx"):
        meta_path = archivo.with_suffix(".meta.yml")
        meta: dict = {}
        if meta_path.exists():
            with meta_path.open(encoding="utf-8") as f:
                meta = yaml.safe_load(f) or {}
        resultado.append({
            "nombre": archivo.stem,
            "ruta": str(archivo.relative_to(REPO_ROOT)).replace("\\", "/"),
            "titulo": meta.get("titulo", archivo.stem),
            "tags": meta.get("tags", []),
        })
    return resultado


def encontrar_plantilla(nombre: str) -> Path | None:
    """Busca una plantilla por nombre exacto (sin extensión) o substring."""
    plantillas_dir = DESPACHO / "plantillas"
    if not plantillas_dir.exists():
        return None
    exacto = list(plantillas_dir.rglob(f"{nombre}.docx"))
    if exacto:
        return exacto[0]
    fuzzy = [p for p in plantillas_dir.rglob("*.docx") if nombre.lower() in p.stem.lower()]
    return fuzzy[0] if fuzzy else None


def encontrar_caso(id_caso: str) -> Path | None:
    """Busca una carpeta de caso por ID exacto o por substring."""
    casos_dir = DESPACHO / "casos"
    if not casos_dir.exists():
        return None
    exacto = casos_dir / id_caso
    if exacto.is_dir():
        return exacto
    candidatos = [d for d in casos_dir.iterdir() if d.is_dir() and id_caso.lower() in d.name.lower()]
    return candidatos[0] if len(candidatos) == 1 else None


# --------- CLI ---------

def _cli() -> int:
    parser = argparse.ArgumentParser(description="Helpers de la skill /editar-documento")
    sub = parser.add_subparsers(dest="cmd", required=True)

    p_slug = sub.add_parser("slug", help="Genera un slug a partir de un texto")
    p_slug.add_argument("texto")

    p_extraer = sub.add_parser("extraer-texto", help="Extrae texto plano de un DOCX")
    p_extraer.add_argument("docx")

    p_listar = sub.add_parser("listar-plantillas", help="Lista plantillas disponibles (JSON)")

    p_buscar_p = sub.add_parser("buscar-plantilla", help="Encuentra plantilla por nombre")
    p_buscar_p.add_argument("nombre")

    p_buscar_c = sub.add_parser("buscar-caso", help="Encuentra carpeta de caso por ID")
    p_buscar_c.add_argument("id")

    p_corr = sub.add_parser("siguiente-correlativo", help="Siguiente NNN del año")
    p_corr.add_argument("anio", type=int)

    args = parser.parse_args()

    if args.cmd == "slug":
        print(slugify(args.texto))
    elif args.cmd == "extraer-texto":
        print(extraer_texto(Path(args.docx)))
    elif args.cmd == "listar-plantillas":
        print(json.dumps(listar_plantillas(), ensure_ascii=False, indent=2))
    elif args.cmd == "buscar-plantilla":
        p = encontrar_plantilla(args.nombre)
        print(str(p.relative_to(REPO_ROOT)).replace("\\", "/") if p else "")
        return 0 if p else 1
    elif args.cmd == "buscar-caso":
        c = encontrar_caso(args.id)
        print(str(c.relative_to(REPO_ROOT)).replace("\\", "/") if c else "")
        return 0 if c else 1
    elif args.cmd == "siguiente-correlativo":
        print(siguiente_correlativo(args.anio, DESPACHO / "casos"))
    return 0


if __name__ == "__main__":
    sys.exit(_cli())
