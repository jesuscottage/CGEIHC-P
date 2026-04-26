"""generar_plantilla_prueba.py — utilidad de soporte para los tests.

Genera una plantilla DOCX representativa con los elementos que la skill
debe manejar correctamente: encabezado, marcadores {{...}}, cláusulas
numeradas, tabla, cita textual entrecomillada, pie con datos del despacho.

Se ejecuta una vez para poblar despacho/plantillas/ con material de prueba.

Uso:
    python generar_plantilla_prueba.py
"""
from __future__ import annotations

import sys
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

REPO_ROOT = Path(__file__).resolve().parents[4]
DESPACHO = REPO_ROOT / "despacho"
PLANTILLAS_DIR = DESPACHO / "plantillas"


def construir_documento() -> Document:
    doc = Document()

    # Estilo base — fuente y tamaño consistentes con un escrito jurídico
    estilo = doc.styles["Normal"]
    estilo.font.name = "Times New Roman"
    estilo.font.size = Pt(12)

    # Encabezado del documento
    seccion = doc.sections[0]
    encabezado = seccion.header.paragraphs[0]
    encabezado.text = "Expediente: {{NUMERO_EXPEDIENTE}} — Fecha: {{FECHA_DOCUMENTO}}"
    encabezado.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Encabezamiento procesal
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("CONTRATO DE ARRENDAMIENTO COMERCIAL")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(
        "C. JUEZ COMPETENTE EN TURNO\nP R E S E N T E"
    ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Comparecencia
    p = doc.add_paragraph()
    p.add_run(
        "{{NOMBRE_ARRENDADOR}}, por mi propio derecho, señalando como domicilio "
        "para oír y recibir notificaciones el ubicado en {{DOMICILIO_ARRENDADOR}}, "
        "comparezco respetuosamente para celebrar el siguiente contrato de "
        "arrendamiento comercial con {{NOMBRE_ARRENDATARIO}}, en los términos "
        "de las cláusulas que a continuación se exponen:"
    )

    # Cláusulas numeradas
    doc.add_paragraph("PRIMERA. Objeto del contrato.").runs[0].bold = True
    doc.add_paragraph(
        "El arrendador concede en arrendamiento al arrendatario el inmueble "
        "ubicado en {{DOMICILIO_INMUEBLE}}, con una superficie aproximada de "
        "{{SUPERFICIE}} metros cuadrados, para uso exclusivamente comercial."
    )

    doc.add_paragraph("SEGUNDA. Plazo.").runs[0].bold = True
    doc.add_paragraph(
        "El presente contrato tendrá una vigencia de {{PLAZO_AÑOS}} años "
        "contados a partir del {{FECHA_INICIO}}, prorrogable por acuerdo "
        "expreso de las partes."
    )

    doc.add_paragraph("TERCERA. Renta.").runs[0].bold = True
    doc.add_paragraph(
        "El arrendatario pagará al arrendador la cantidad de {{RENTA_MENSUAL}} "
        "pesos mensuales, dentro de los primeros cinco días de cada mes, "
        "mediante depósito o transferencia a la cuenta {{CUENTA_BANCARIA}}."
    )

    doc.add_paragraph("CUARTA. Garantía.").runs[0].bold = True
    doc.add_paragraph(
        "Como garantía de cumplimiento, el arrendatario entregará un depósito "
        "equivalente a un mes de renta, el cual será devuelto al término del "
        "contrato siempre que no existan adeudos ni daños al inmueble."
    )

    doc.add_paragraph("QUINTA. Causales de rescisión.").runs[0].bold = True
    doc.add_paragraph(
        "Cualquiera de las partes podrá rescindir el contrato por las causales "
        "previstas en el Código Civil aplicable. La parte afectada deberá "
        "notificar a la contraparte con al menos treinta días de anticipación."
    )

    # Tabla — sumario económico
    doc.add_paragraph()
    doc.add_paragraph("SUMARIO ECONÓMICO").runs[0].bold = True
    tabla = doc.add_table(rows=4, cols=2)
    tabla.style = "Table Grid"
    tabla.cell(0, 0).text = "Concepto"
    tabla.cell(0, 1).text = "Monto"
    tabla.cell(1, 0).text = "Renta mensual"
    tabla.cell(1, 1).text = "{{RENTA_MENSUAL}}"
    tabla.cell(2, 0).text = "Depósito en garantía"
    tabla.cell(2, 1).text = "{{DEPOSITO}}"
    tabla.cell(3, 0).text = "Plazo total"
    tabla.cell(3, 1).text = "{{PLAZO_AÑOS}} años"

    # Fundamento legal — cita textual entrecomillada (NO se debe alterar)
    doc.add_paragraph()
    doc.add_paragraph("FUNDAMENTO LEGAL").runs[0].bold = True
    cita = doc.add_paragraph()
    cita.add_run(
        '"El arrendamiento es un contrato por el cual una parte llamada '
        "arrendador concede a otra llamada arrendatario el uso o goce "
        "temporal de una cosa, mediante el pago de un precio cierto.\" "
        "(Código Civil Federal, artículo 2398)."
    ).italic = True

    # Pie de firma del abogado — convención intocable
    doc.add_paragraph()
    doc.add_paragraph()
    firma = doc.add_paragraph()
    firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firma.add_run("_______________________________________\n").bold = False
    firma.add_run("Lic. Juan Pablo (autor de la plantilla)\n").bold = True
    firma.add_run("Cédula profesional: 1234567\n")
    firma.add_run("Domicilio del despacho: Calle Ejemplo 123, Ciudad")

    return doc


def main() -> int:
    PLANTILLAS_DIR.mkdir(parents=True, exist_ok=True)
    salida = PLANTILLAS_DIR / "ejemplo-plantilla-v1.docx"
    meta_salida = PLANTILLAS_DIR / "ejemplo-plantilla-v1.meta.yml"

    doc = construir_documento()
    doc.save(str(salida))

    meta = {
        "titulo": "Contrato de arrendamiento comercial — plantilla de prueba",
        "autor": "Juan Pablo",
        "fecha_creacion": "2026-04-25",
        "fecha_revision": "2026-04-25",
        "variables_detectadas": [
            "NUMERO_EXPEDIENTE",
            "FECHA_DOCUMENTO",
            "NOMBRE_ARRENDADOR",
            "DOMICILIO_ARRENDADOR",
            "NOMBRE_ARRENDATARIO",
            "DOMICILIO_INMUEBLE",
            "SUPERFICIE",
            "PLAZO_AÑOS",
            "FECHA_INICIO",
            "RENTA_MENSUAL",
            "CUENTA_BANCARIA",
            "DEPOSITO",
        ],
        "tags": ["contrato", "arrendamiento", "comercial", "plantilla-prueba"],
        "uso_historico": {
            "casos_exitosos": 0,
            "ultimo_uso": None,
            "notas": "Plantilla generada para los tests integrados de la skill /editar-documento.",
        },
        "contiene_pii": False,
    }
    with meta_salida.open("w", encoding="utf-8") as f:
        yaml.safe_dump(meta, f, allow_unicode=True, sort_keys=False, indent=2, default_flow_style=False)

    print(f"OK plantilla generada: {salida.relative_to(REPO_ROOT)}")
    print(f"OK meta generada:      {meta_salida.relative_to(REPO_ROOT)}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
